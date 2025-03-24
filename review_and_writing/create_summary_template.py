#!/usr/bin/env python
"""
Combined Summary Template Generator

This script generates a Word document template for the combined summary of
all articles analyzed in the systematic review. It creates a structured document
with placeholders that can be filled in manually or programmatically.

Usage:
    python create_summary_template.py [--output OUTPUT_FILE]

Options:
    --output OUTPUT_FILE  Name of the output file (default: combined_summary_template.docx)
"""

import os
import argparse
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from datetime import datetime

def add_conclusion_section(doc):
    """Add a conclusions section to the document."""
    doc.add_heading("4. Conclusions", level=1)
    
    conclusions = doc.add_paragraph("", style='Placeholder')
    conclusions.add_run("Provide a concise summary of the main conclusions from this systematic review, including:")
    conclusions.add_run("\n• Primary takeaways regarding effective employment interventions for people with mental health disorders")
    conclusions.add_run("\n• Level of confidence in these conclusions based on the evidence")
    conclusions.add_run("\n• Key recommendations for different stakeholders")
    conclusions.add_run("\n• Areas requiring further research or development")
    
    doc.add_page_break()

def add_references_section(doc):
    """Add a references section to the document."""
    doc.add_heading("5. References", level=1)
    
    references = doc.add_paragraph("", style='Placeholder')
    references.add_run("[Insert references in appropriate citation format]")
    references.add_run("\n\nNote: References should include all cited works as well as the included studies.")
    
    doc.add_page_break()

def add_appendices(doc):
    """Add appendices to the document."""
    doc.add_heading("Appendices", level=1)
    
    # Appendix A: Included Studies
    doc.add_heading("Appendix A: Included Studies", level=2)
    appendix_a = doc.add_paragraph("", style='Placeholder')
    appendix_a.add_run("Provide a detailed table of all included studies with the following information:")
    appendix_a.add_run("\n• Article ID and citation")
    appendix_a.add_run("\n• Study design and quality assessment")
    appendix_a.add_run("\n• Population characteristics")
    appendix_a.add_run("\n• Intervention details")
    appendix_a.add_run("\n• Key outcomes")
    appendix_a.add_run("\n• Summary of findings")
    
    # Appendix B: Search Strategy
    doc.add_heading("Appendix B: Search Strategy", level=2)
    appendix_b = doc.add_paragraph("", style='Placeholder')
    appendix_b.add_run("Document the complete search strategy used for the systematic review, including:")
    appendix_b.add_run("\n• Databases searched")
    appendix_b.add_run("\n• Search terms and combinations")
    appendix_b.add_run("\n• Search date")
    appendix_b.add_run("\n• Inclusion and exclusion criteria details")
    appendix_b.add_run("\n• PRISMA flow diagram")
    
    # Appendix C: LLM Analysis Methods
    doc.add_heading("Appendix C: LLM Analysis Methods", level=2)
    appendix_c = doc.add_paragraph("", style='Placeholder')
    appendix_c.add_run("Provide details on the use of LLMs for article analysis, including:")
    appendix_c.add_run("\n• Models used and their specifications")
    appendix_c.add_run("\n• Prompts for each section (background, methods, results, etc.)")
    appendix_c.add_run("\n• Quality control measures")
    appendix_c.add_run("\n• Limitations of the approach")

def parse_arguments():
    """Parse command line arguments."""
    parser = argparse.ArgumentParser(description='Generate a template for the combined summary document')
    parser.add_argument('--output', type=str, default='combined_summary_template.docx',
                        help='Name of the output file')
    return parser.parse_args()

def setup_document_styles(doc):
    """Set up custom styles for the document."""
    # Title style
    title_style = doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.base_style = doc.styles['Title']
    title_style.font.name = 'Calibri'
    title_style.font.size = Pt(24)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0, 70, 127)
    title_paragraph_format = title_style.paragraph_format
    title_paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph_format.space_after = Pt(12)
    
    # Heading 1 style
    h1_style = doc.styles['Heading 1']
    h1_style.font.name = 'Calibri'
    h1_style.font.size = Pt(16)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(0, 70, 127)
    h1_paragraph_format = h1_style.paragraph_format
    h1_paragraph_format.space_before = Pt(24)
    h1_paragraph_format.space_after = Pt(12)
    
    # Heading 2 style
    h2_style = doc.styles['Heading 2']
    h2_style.font.name = 'Calibri'
    h2_style.font.size = Pt(14)
    h2_style.font.bold = True
    h2_style.font.color.rgb = RGBColor(0, 90, 156)
    h2_paragraph_format = h2_style.paragraph_format
    h2_paragraph_format.space_before = Pt(18)
    h2_paragraph_format.space_after = Pt(6)
    
    # Normal text style
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_paragraph_format = normal_style.paragraph_format
    normal_paragraph_format.space_after = Pt(8)
    
    # Custom placeholder style
    placeholder_style = doc.styles.add_style('Placeholder', WD_STYLE_TYPE.PARAGRAPH)
    placeholder_style.base_style = doc.styles['Normal']
    placeholder_style.font.name = 'Calibri'
    placeholder_style.font.size = Pt(11)
    placeholder_style.font.italic = True
    placeholder_style.font.color.rgb = RGBColor(127, 127, 127)
    
    # Table styles
    table_style = doc.styles.add_style('CustomTable', WD_STYLE_TYPE.TABLE)
    table_style.base_style = doc.styles['Table Grid']
    
    return doc

def add_cover_page(doc):
    """Add a cover page to the document."""
    # Add title
    title = doc.add_paragraph("Systematic Review of Employment Interventions for People with Mental Health Disorders", style='CustomTitle')
    
    # Add subtitle
    subtitle = doc.add_paragraph("Combined Analysis and Synthesis of Included Studies")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(18)
    subtitle.runs[0].font.italic = True
    subtitle.runs[0].font.color.rgb = RGBColor(0, 90, 156)
    
    # Add date
    date = doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date.runs[0].font.size = Pt(12)
    
    # Add authors placeholder
    authors = doc.add_paragraph("Authors: [Insert Author Names]")
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    authors.runs[0].font.size = Pt(12)
    
    # Add page break
    doc.add_page_break()

def add_table_of_contents(doc):
    """Add a table of contents page to the document."""
    # Add TOC heading
    doc.add_heading("Table of Contents", level=1)
    
    # Add placeholder text for TOC
    toc = doc.add_paragraph("", style='Placeholder')
    run = toc.add_run("Note: This Table of Contents will be populated automatically when the document is finalized. Right-click and select 'Update Field' to update it.")
    run.italic = True
    run.font.color.rgb = RGBColor(127, 127, 127)
    
    # Add page break
    doc.add_page_break()

def add_executive_summary(doc):
    """Add an executive summary section to the document."""
    doc.add_heading("Executive Summary", level=1)
    
    summary = doc.add_paragraph("", style='Placeholder')
    summary.add_run("This section should provide a brief overview of the systematic review, including:")
    summary.add_run("\n• The research question addressed")
    summary.add_run("\n• The scope and methodology of the review")
    summary.add_run("\n• Key findings across all included studies")
    summary.add_run("\n• Main conclusions and implications for practice")
    summary.add_run("\n\nRecommended length: 500-750 words")
    
    doc.add_page_break()

def add_introduction_section(doc):
    """Add an introduction section to the document."""
    doc.add_heading("1. Introduction", level=1)
    
    # Background subsection
    doc.add_heading("1.1 Background", level=2)
    background = doc.add_paragraph("", style='Placeholder')
    background.add_run("Provide context on employment interventions for people with mental health disorders. Discuss:")
    background.add_run("\n• The impact of mental health disorders on employment")
    background.add_run("\n• Current challenges in employment support")
    background.add_run("\n• Types of interventions commonly used")
    background.add_run("\n• The need for evidence synthesis in this area")
    
    # Research question subsection
    doc.add_heading("1.2 Research Question", level=2)
    question = doc.add_paragraph("", style='Placeholder')
    question.add_run("State the primary research question(s) and objectives of the review. Example:")
    question.add_run("\n\n\"This systematic review aims to identify what works in active employment policies for people with mental health disorders, focusing on evidence from EU countries, Europe, US, Canada, and Australia published between 2005 and 2025.\"")
    
    # Methods overview subsection
    doc.add_heading("1.3 Methods Overview", level=2)
    methods = doc.add_paragraph("", style='Placeholder')
    methods.add_run("Briefly describe the methodology used for the systematic review, including:")
    methods.add_run("\n• Inclusion and exclusion criteria")
    methods.add_run("\n• Search strategy")
    methods.add_run("\n• Article screening and selection process")
    methods.add_run("\n• Analysis approach using LLMs")
    methods.add_run("\n\nMore detailed methodology can be included in a separate section or appendix.")
    
    doc.add_page_break()

def add_findings_section(doc):
    """Add a findings section to the document."""
    doc.add_heading("2. Synthesis of Findings", level=1)
    
    # Overview of included studies
    doc.add_heading("2.1 Overview of Included Studies", level=2)
    overview = doc.add_paragraph("", style='Placeholder')
    overview.add_run("Provide a summary of the included studies, including:")
    overview.add_run("\n• Number of studies by type (RCT, quasi-experimental, etc.)")
    overview.add_run("\n• Distribution by geographic region")
    overview.add_run("\n• Range of sample sizes")
    overview.add_run("\n• Types of mental health conditions studied")
    overview.add_run("\n• Types of employment interventions evaluated")
    
    # Add a sample table for study characteristics
    doc.add_paragraph("Table 1: Summary of Included Studies")
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    
    # Set header row
    header_cells = table.rows[0].cells
    header_cells[0].text = "Study ID"
    header_cells[1].text = "Study Design"
    header_cells[2].text = "Sample Size"
    header_cells[3].text = "Mental Health Condition"
    header_cells[4].text = "Intervention Type"
    
    # Add a few placeholder rows
    for i in range(3):
        row_cells = table.add_row().cells
        for j in range(5):
            row_cells[j].text = "[Insert data]"
    
    doc.add_paragraph("Note: Full details of all included studies are available in Appendix A.", style='Placeholder')
    
    # Effectiveness by intervention type
    doc.add_heading("2.2 Effectiveness by Intervention Type", level=2)
    effectiveness = doc.add_paragraph("", style='Placeholder')
    effectiveness.add_run("Analyze the effectiveness of different intervention types (group findings by intervention category). For each type, discuss:")
    effectiveness.add_run("\n• Overall effectiveness for employment outcomes")
    effectiveness.add_run("\n• Variations in implementation and context")
    effectiveness.add_run("\n• Factors that may influence effectiveness")
    effectiveness.add_run("\n• Quality and strength of the evidence")
    
    # Subsections for common intervention types
    intervention_types = [
        "2.2.1 Supported Employment",
        "2.2.2 Vocational Rehabilitation",
        "2.2.3 Skills Training",
        "2.2.4 Job Search Assistance",
        "2.2.5 Workplace Accommodations",
        "2.2.6 Other Intervention Types"
    ]
    
    for subsection in intervention_types:
        doc.add_heading(subsection, level=3)
        p = doc.add_paragraph("", style='Placeholder')
        p.add_run("[Insert analysis of effectiveness for this intervention type]")
    
    # Effectiveness by mental health condition
    doc.add_heading("2.3 Effectiveness by Mental Health Condition", level=2)
    by_condition = doc.add_paragraph("", style='Placeholder')
    by_condition.add_run("Analyze how intervention effectiveness varies by mental health condition. Consider including subsections for:")
    by_condition.add_run("\n• Depression")
    by_condition.add_run("\n• Anxiety disorders")
    by_condition.add_run("\n• Schizophrenia and psychotic disorders")
    by_condition.add_run("\n• Bipolar disorder")
    by_condition.add_run("\n• Personality disorders")
    by_condition.add_run("\n• General mental health/mixed conditions")
    
    # Implementation considerations
    doc.add_heading("2.4 Implementation Considerations", level=2)
    implementation = doc.add_paragraph("", style='Placeholder')
    implementation.add_run("Discuss factors that affect the implementation and success of interventions:")
    implementation.add_run("\n• Integration with mental health services")
    implementation.add_run("\n• Service provider training and qualifications")
    implementation.add_run("\n• Employer engagement strategies")
    implementation.add_run("\n• Duration and intensity of support")
    implementation.add_run("\n• Sustainability and cost considerations")
    
    doc.add_page_break()

def add_discussion_section(doc):
    """Add a discussion section to the document."""
    doc.add_heading("3. Discussion", level=1)
    
    # Key findings
    doc.add_heading("3.1 Summary of Key Findings", level=2)
    key_findings = doc.add_paragraph("", style='Placeholder')
    key_findings.add_run("Summarize the most important findings of the systematic review, emphasizing:")
    key_findings.add_run("\n• Which interventions show the strongest evidence of effectiveness")
    key_findings.add_run("\n• Common elements of successful programs")
    key_findings.add_run("\n• Variations in effectiveness by context or population")
    key_findings.add_run("\n• Unexpected or contrasting findings")
    
    # Comparison with existing literature
    doc.add_heading("3.2 Comparison with Existing Literature", level=2)
    comparison = doc.add_paragraph("", style='Placeholder')
    comparison.add_run("Discuss how the findings of this review compare with:")
    comparison.add_run("\n• Previous systematic reviews on similar topics")
    comparison.add_run("\n• Established theories and models in vocational rehabilitation")
    comparison.add_run("\n• Current best practice guidelines")
    comparison.add_run("\n• Emerging trends in employment support")
    
    # Strengths and limitations
    doc.add_heading("3.3 Strengths and Limitations", level=2)
    strengths_limitations = doc.add_paragraph("", style='Placeholder')
    strengths_limitations.add_run("Discuss the strengths and limitations of:")
    strengths_limitations.add_run("\n• The included studies (methodological quality, sample sizes, etc.)")
    strengths_limitations.add_run("\n• This systematic review process (scope, search strategy, analysis method)")
    strengths_limitations.add_run("\n• The use of LLMs for article analysis and synthesis")
    strengths_limitations.add_run("\n• Generalizability of the findings")
    
    # Implications
    doc.add_heading("3.4 Implications for Practice and Policy", level=2)
    implications = doc.add_paragraph("", style='Placeholder')
    implications.add_run("Discuss the practical implications of the findings for:")
    implications.add_run("\n• Employment support providers")
    implications.add_run("\n• Mental health practitioners")
    implications.add_run("\n• Employers and workplace managers")
    implications.add_run("\n• Policymakers and funders")
    implications.add_run("\n• Individuals with mental health disorders seeking employment")
    
    doc.add_page_break()

def main():
    """Main function to generate the template document."""
    args = parse_arguments()
    
    # Create a new document
    doc = docx.Document()
    
    # Set up document styles
    doc = setup_document_styles(doc)
    
    # Add document components in order
    add_cover_page(doc)
    add_table_of_contents(doc)
    add_executive_summary(doc)
    add_introduction_section(doc)
    add_findings_section(doc)
    add_discussion_section(doc)
    add_conclusion_section(doc)
    add_references_section(doc)
    add_appendices(doc)
    
    # Save the document
    doc.save(args.output)
    print(f"Template document created and saved to: {args.output}")

if __name__ == "__main__":
    main()