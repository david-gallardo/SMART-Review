#!/usr/bin/env python
"""
Chunked Batch PDF Article Processor

This script divides a large collection of PDF articles into smaller batches
and processes each batch sequentially using the batch_processor.py script.

Usage:
    python batch_processor_chunked.py [--batch_size BATCH_SIZE] [--parallel PARALLEL]
                                     [--model MODEL_NAME] [--api_url API_URL]

Options:
    --batch_size BATCH_SIZE  Number of articles in each batch (default: 10)
    --parallel PARALLEL      Number of articles to process in parallel within each batch (default: 4)
    --model MODEL_NAME       Specify the LLM model to use
    --api_url API_URL        Custom LLM API URL if needed
"""

import os
import sys
import argparse
import shutil
import subprocess
import time
from datetime import datetime
import logging

# Set up logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("batch_processor_chunked.log"),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

def parse_arguments():
    """Parse command line arguments."""
    parser = argparse.ArgumentParser(description='Process PDFs in smaller batches')
    parser.add_argument('--batch_size', type=int, default=10,
                        help='Number of articles in each batch')
    parser.add_argument('--parallel', type=int, default=4,
                        help='Number of articles to process in parallel within each batch')
    parser.add_argument('--model', type=str, default='mistral-small-24b-instruct-2501',
                        help='LLM model to use for analysis')
    parser.add_argument('--api_url', type=str, help='Custom LLM API URL')
    return parser.parse_args()

def setup_directories():
    """Set up the necessary directories for batch processing."""
    # Get the current directory as the base directory
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    
    # Define directory paths
    articles_dir = os.path.join(base_dir, 'documents', 'articles')
    temp_dir = os.path.join(base_dir, 'documents', 'temp_processing')
    output_dir = os.path.join(base_dir, 'output', 'summaries')
    
    # Create temp directory if it doesn't exist
    if not os.path.exists(temp_dir):
        os.makedirs(temp_dir)
    
    # Ensure output directory exists
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    return base_dir, articles_dir, temp_dir, output_dir

def get_pdf_files(directory):
    """Get all PDF files in the specified directory."""
    return [f for f in os.listdir(directory) if f.lower().endswith('.pdf')]

def process_batch(batch_files, source_dir, temp_dir, args, batch_num, total_batches):
    """Process a batch of PDF files."""
    batch_start_time = time.time()
    logger.info(f"Starting Batch {batch_num}/{total_batches} - Processing {len(batch_files)} files")
    
    # Move files to temp directory
    for filename in batch_files:
        source_path = os.path.join(source_dir, filename)
        dest_path = os.path.join(temp_dir, filename)
        shutil.copy2(source_path, dest_path)
    
    # Build command for batch_processor.py
    cmd = [
        sys.executable, "batch_processor.py",
        "--parallel", str(args.parallel),
        "--output", f"batch_{batch_num}_summary.docx",
    ]
    
    # Add optional arguments if provided
    if args.model:
        cmd.extend(["--model", args.model])
    if args.api_url:
        cmd.extend(["--api_url", args.api_url])
    
    # Run batch_processor.py
    try:
        logger.info(f"Running command: {' '.join(cmd)}")
        process = subprocess.Popen(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, universal_newlines=True)
        
        # Monitor output in real-time
        while True:
            output = process.stdout.readline()
            if output == '' and process.poll() is not None:
                break
            if output:
                logger.info(output.strip())
        
        return_code = process.wait()
        
        if return_code != 0:
            stderr = process.stderr.read()
            logger.error(f"Error processing batch {batch_num}: {stderr}")
            return False
        
    except Exception as e:
        logger.error(f"Exception during processing of batch {batch_num}: {str(e)}")
        return False
    finally:
        # Clean up temp directory
        for filename in os.listdir(temp_dir):
            file_path = os.path.join(temp_dir, filename)
            if os.path.isfile(file_path):
                os.remove(file_path)
    
    batch_duration = time.time() - batch_start_time
    logger.info(f"Completed Batch {batch_num}/{total_batches} in {batch_duration/60:.2f} minutes")
    return True

def combine_batch_summaries(output_dir, total_batches):
    """Create a final combined document from all batch summaries."""
    # This is a placeholder for more sophisticated combination logic
    # In a real implementation, you might want to use python-docx to combine the documents
    
    logger.info("Creating combined summary from all batches")
    combined_file = os.path.join(output_dir, "all_articles_combined_summary.docx")
    
    # For now, we'll just create a simple document listing all batch summaries
    try:
        import docx
        doc = docx.Document()
        
        doc.add_heading("Combined Article Summaries", level=0)
        doc.add_paragraph(f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        doc.add_paragraph("This document is an index of all batch summaries. Please refer to the individual batch summary documents for detailed analysis.")
        
        for i in range(1, total_batches + 1):
            batch_file = f"batch_{i}_summary.docx"
            doc.add_paragraph(f"Batch {i}: {batch_file}", style='ListBullet')
        
        doc.save(combined_file)
        logger.info(f"Created combined summary index at {combined_file}")
        return True
    except Exception as e:
        logger.error(f"Error creating combined summary: {str(e)}")
        return False

def main():
    """Main function to process PDFs in batches."""
    start_time = time.time()
    args = parse_arguments()
    base_dir, articles_dir, temp_dir, output_dir = setup_directories()
    
    # Get all PDF files
    all_files = get_pdf_files(articles_dir)
    if not all_files:
        logger.error("No PDF files found in the articles directory")
        return
    
    total_files = len(all_files)
    logger.info(f"Found {total_files} PDF files to process")
    
    # Divide into batches
    batch_size = min(args.batch_size, total_files)
    batches = [all_files[i:i + batch_size] for i in range(0, total_files, batch_size)]
    total_batches = len(batches)
    
    logger.info(f"Divided into {total_batches} batches of {batch_size} files each")
    
    # Process each batch
    successful_batches = 0
    for i, batch in enumerate(batches, 1):
        success = process_batch(batch, articles_dir, temp_dir, args, i, total_batches)
        if success:
            successful_batches += 1
        
        # Optional: Add a short pause between batches
        if i < total_batches:
            logger.info(f"Pausing for 5 seconds before starting next batch")
            time.sleep(5)
    
    # Combine results from all batches
    if successful_batches > 0:
        combine_batch_summaries(output_dir, total_batches)
    
    # Clean up temp directory
    shutil.rmtree(temp_dir)
    os.makedirs(temp_dir)
    
    total_duration = time.time() - start_time
    logger.info(f"Processing complete! Processed {successful_batches} of {total_batches} batches in {total_duration/60:.2f} minutes")
    logger.info(f"Individual summaries and batch summaries saved to {output_dir}")

if __name__ == "__main__":
    main()