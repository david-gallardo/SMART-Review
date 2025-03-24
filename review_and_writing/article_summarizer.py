#!/usr/bin/env python
"""
PDF Article Processor and Summarizer

This script:
1. Reads PDF articles from the 'documents/articles' directory
2. Extracts the text content from each PDF
3. Sends structured prompts to LLMs based on your predefined section templates
4. Compiles the responses into a structured Word document

Usage:
    python article_summarizer.py [--test] [--article_id ARTICLE_ID]

Options:
    --test              Process only one PDF file for testing
    --article_id ID     Process only the article with the specified ID

Requirements:
    - pandas
    - PyPDF2 (for PDF extraction)
    - python-docx (for Word document creation)
    - requests (for LLM API calls)
"""

import os
import sys
import argparse
import json
import re
import pandas as pd
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import requests
import concurrent.futures
import time
from datetime import datetime

# Import the improved PDF extraction utilities
from pdf_extraction_utils import (
    improve_pdf_extraction, 
    clean_scientific_text,
    extract_with_section_detection,
    detect_article_structure
)

# Import the PromptManager for advanced prompt handling
from advanced_prompt_manager import PromptManager

def parse_arguments():
    """Parse command line arguments."""
    parser = argparse.ArgumentParser(description='Process and summarize PDF articles using LLMs')
    parser.add_argument('--test', action='store_true', help='Process only one PDF file for testing')
    parser.add_argument('--article_id', type=int, help='Process only the article with the specified ID')
    parser.add_argument('--api_url', type=str, help='Custom LLM API URL')
    parser.add_argument('--prompts_dir', type=str, help='Directory containing custom prompt templates')
    parser.add_argument('--articles_dir', type=str, help='Directory containing PDF articles')
    parser.add_argument('--output_dir', type=str, help='Directory for output files')
    parser.add_argument('--model', type=str, default='mistral-small-24b-instruct-2501', 
                      help='LLM model to use (default: mistral-small-24b-instruct-2501)')
    return parser.parse_args()

def setup_directories(args):
    """Set up the necessary directories for the script."""
    # Get the current directory as the base directory
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    
    # Use command-line arguments if provided, otherwise use default paths
    data_dir = getattr(args, 'data_dir', None) or os.path.join(base_dir, 'data')
    articles_dir = getattr(args, 'articles_dir', None) or os.path.join(base_dir, 'documents', 'articles')
    output_dir = getattr(args, 'output_dir', None) or os.path.join(base_dir, 'output', 'summaries')
    prompts_dir = getattr(args, 'prompts_dir', None) or os.path.join(base_dir, 'prompts')
    
    # Create directories if they don't exist
    os.makedirs(data_dir, exist_ok=True)
    os.makedirs(articles_dir, exist_ok=True)
    os.makedirs(output_dir, exist_ok=True)
    os.makedirs(prompts_dir, exist_ok=True)
    
    return base_dir, data_dir, articles_dir, output_dir, prompts_dir

def load_article_metadata(data_dir):
    """Load metadata for the included articles."""
    try:
        # Try to load from the classified file first
        classified_file = os.path.join(data_dir, 'df_articles_results_classified.xlsx')
        if os.path.exists(classified_file):
            df = pd.read_excel(classified_file)
            included_df = df[df['GlobalInclusion'] == 'Yes'].copy()
            print(f"Loaded {len(included_df)} included articles from classification data.")
            return included_df
        
        # If not found, try the basic results file
        results_file = os.path.join(data_dir, 'df_articles_results.xlsx')
        if os.path.exists(results_file):
            df = pd.read_excel(results_file)
            included_df = df[df['GlobalInclusion'] == 'Yes'].copy()
            print(f"Loaded {len(included_df)} included articles from results data.")
            return included_df
        
        print("No article metadata found. Will process PDF files without metadata.")
        return None
    
    except Exception as e:
        print(f"Error loading article metadata: {e}")
        return None

def get_pdf_files(articles_dir, article_id=None):
    """Get list of PDF files to process, optionally filtered by article ID."""
    pdf_files = []
    
    if not os.path.exists(articles_dir):
        print(f"Articles directory not found: {articles_dir}")
        return pdf_files
    
    # If article_id is provided, look for a specific pattern
    if article_id is not None:
        for filename in os.listdir(articles_dir):
            if filename.lower().endswith('.pdf'):
                # Look for article ID in the filename (various formats)
                patterns = [
                    fr'^{article_id}_',  # starts with ID_
                    fr'_{article_id}_',  # contains _ID_
                    fr'_{article_id}\.pdf$',  # ends with _ID.pdf
                    fr'^article_{article_id}',  # starts with article_ID
                ]
                
                if any(re.search(pattern, filename) for pattern in patterns):
                    pdf_files.append(os.path.join(articles_dir, filename))
                    break
    else:
        # Get all PDF files
        pdf_files = [os.path.join(articles_dir, f) for f in os.listdir(articles_dir) 
                   if f.lower().endswith('.pdf')]
    
    return pdf_files

def chunk_text(text, max_chunk_length=4000, overlap=500):
    """
    Split the text into overlapping chunks for LLM processing.
    
    Args:
        text (str): The text to chunk
        max_chunk_length (int): Maximum length for each chunk
        overlap (int): How much text should overlap between chunks
        
    Returns:
        list: List of text chunks
    """
    if not text:
        return []
    
    # If text is shorter than max chunk length, return as a single chunk
    if len(text) <= max_chunk_length:
        return [text]
    
    chunks = []
    start = 0
    
    while start < len(text):
        # Calculate end position
        end = min(start + max_chunk_length, len(text))
        
        # If we're not at the end of the text, try to find a good break point
        if end < len(text):
            # Look for paragraph break
            paragraph_break = text.rfind('\n\n', start, end)
            if paragraph_break != -1 and paragraph_break > start + max_chunk_length // 2:
                end = paragraph_break
            else:
                # Look for sentence break
                sentence_break = text.rfind('. ', start, end)
                if sentence_break != -1 and sentence_break > start + max_chunk_length // 2:
                    end = sentence_break + 1  # Include the period
                else:
                    # Last resort: word break
                    word_break = text.rfind(' ', start, end)
                    if word_break != -1 and word_break > start + max_chunk_length // 2:
                        end = word_break
        
        chunks.append(text[start:end].strip())
        
        # Move start position for next chunk, with overlap
        start = max(start + max_chunk_length - overlap, end - overlap)
    
    return chunks

def query_llm(model_name, prompt, text_chunk, api_url=None):
    """
    Query the LLM with a prompt and article text.
    
    Args:
        model_name (str): The name of the LLM model to use
        prompt (str): The prompt to send to the LLM
        text_chunk (str): The article text chunk to analyze
        api_url (str, optional): Custom API URL
        
    Returns:
        str: The LLM's response
    """
    url = api_url or "http://127.0.0.1:1234/v1/chat/completions"
    
    # Shorten the text_chunk if it's too long
    if len(text_chunk) > 8000:
        text_chunk = text_chunk[:8000] + "..."
    
    # Simplify the system message
    combined_prompt = f"{prompt}\n\nHere is the article text:\n\n{text_chunk}"
    
    request_body = {
        "model": model_name,
        "messages": [
            {"role": "system", "content": "You are an expert research assistant."},
            {"role": "user", "content": combined_prompt}
        ],
        "temperature": 0.3,
        "max_tokens": 1000
    }
    
    try:
        response = requests.post(url, json=request_body, headers={"Content-Type": "application/json"})
        response.raise_for_status()
        
        response_json = response.json()
        generated_text = response_json["choices"][0]["message"]["content"]
        
        return generated_text.strip()
    
    except Exception as e:
        print(f"Error querying LLM ({model_name}): {e}")
        # Add a retry mechanism
        print("Retrying with a shorter text chunk...")
        if len(text_chunk) > 4000:
            # If error occurs, retry with a shorter chunk
            return query_llm(model_name, prompt, text_chunk[:4000] + "...", api_url)
        return f"Error processing this section: {str(e)}"

def process_all_sections(model_name, article_text, prompt_manager, article_info=None, article_id=None, filename=None, api_url=None):
    """
    Process all sections of the article text using the LLM.
    
    Args:
        model_name (str): The name of the LLM model to use
        article_text (str): The complete article text
        prompt_manager (PromptManager): The prompt manager instance
        article_info (dict, optional): Information about the article for specialized prompts
        article_id (int, optional): ID of the article
        filename (str, optional): Filename of the article
        api_url (str, optional): Custom API URL
        
    Returns:
        dict: Results for each section
    """
    results = {
        "article_id": article_id if article_id is not None else "unknown",
        "filename": filename if filename is not None else "unknown",
        "processing_time": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "sections": {}
    }
    
    # Split the text into chunks for processing
    text_chunks = chunk_text(article_text)
    
    if not text_chunks:
        print(f"Warning: No valid text chunks created for article ID {article_id}")
        return results
    
    print(f"Processing article{' ID '+str(article_id) if article_id else ''} with {len(text_chunks)} text chunks")
    
    # Get all available sections from the prompt manager
    section_names = prompt_manager.base_prompts.keys()
    
    # Process each section
    for section_name in section_names:
        print(f"  Processing section: {section_name}")
        
        # Get the appropriate prompt from the prompt manager
        section_prompt = prompt_manager.get_prompt(section_name, article_info)
        
        # For the synthesis section, we'll use all chunks to provide a comprehensive view
        if section_name == "synthesis":
            # If there are too many chunks, use a sample of them
            if len(text_chunks) > 3:
                # Use first, middle and last chunk for synthesis
                selected_chunks = [
                    text_chunks[0],
                    text_chunks[len(text_chunks)//2],
                    text_chunks[-1]
                ]
                combined_text = "\n\n[...]\n\n".join(selected_chunks)
            else:
                combined_text = "\n\n".join(text_chunks)
                
            section_result = query_llm(model_name, section_prompt, combined_text, api_url)
            time.sleep(2)
        
        # For other sections, we'll use the most relevant chunks
        # This is a simple heuristic - could be improved with better section detection
        elif section_name in ["background", "conclusions"]:
            # For background and conclusions, prioritize first and last chunks
            if len(text_chunks) > 1:
                section_result = query_llm(model_name, section_prompt, 
                                          text_chunks[0] + "\n\n[...]\n\n" + text_chunks[-1], api_url)
            else:
                section_result = query_llm(model_name, section_prompt, text_chunks[0], api_url)
                
        elif section_name == "methods":
            # Methods often appear early in the paper
            if len(text_chunks) > 1:
                # Use first and second chunks
                section_text = text_chunks[0]
                if len(text_chunks) > 1:
                    section_text += "\n\n" + text_chunks[1]
                section_result = query_llm(model_name, section_prompt, section_text, api_url)
            else:
                section_result = query_llm(model_name, section_prompt, text_chunks[0], api_url)
                
        elif section_name in ["results", "discussion"]:
            # Results and discussion typically appear later
            if len(text_chunks) > 2:
                # Use middle chunks
                mid_index = len(text_chunks) // 2
                section_text = text_chunks[mid_index]
                if mid_index + 1 < len(text_chunks):
                    section_text += "\n\n" + text_chunks[mid_index + 1]
                section_result = query_llm(model_name, section_prompt, section_text, api_url)
            else:
                # If only 1-2 chunks, use the last available
                section_result = query_llm(model_name, section_prompt, text_chunks[-1], api_url)
        
        else:
            # Default: use the entire text (combined chunks)
            combined_text = "\n\n".join(text_chunks)
            section_result = query_llm(model_name, section_prompt, combined_text, api_url)
        
        results["sections"][section_name] = section_result
    
    return results

def create_word_document(results, article_metadata=None, output_path=None):
    """
    Create a Word document with the article summary.
    
    Args:
        results (dict): The results from processing the article
        article_metadata (pd.Series, optional): Metadata for the article
        output_path (str, optional): Path to save the document
        
    Returns:
        docx.Document: The created Word document
    """
    doc = docx.Document()
    
    # Set up document styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # Add title
    title_text = "Article Summary"
    if article_metadata is not None and 'Article Title' in article_metadata:
        title_text = article_metadata['Article Title']
    elif results.get('filename'):
        title_text = f"Summary of {os.path.basename(results['filename'])}"
    
    title = doc.add_heading(title_text, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add metadata section
    doc.add_heading("Article Metadata", level=2)
    
    metadata_table = doc.add_table(rows=1, cols=2)
    metadata_table.style = 'Table Grid'
    
    # Set column widths
    for cell in metadata_table.columns[0].cells:
        cell.width = Inches(1.5)
    for cell in metadata_table.columns[1].cells:
        cell.width = Inches(4.5)
    
    # Add metadata fields
    hdr_cells = metadata_table.rows[0].cells
    hdr_cells[0].text = "Field"
    hdr_cells[1].text = "Value"
    
    metadata_fields = [
        ("Article ID", results.get("article_id", "N/A")),
        ("Filename", os.path.basename(results.get("filename", "N/A"))),
        ("Processing Date", results.get("processing_time", "N/A")),
    ]
    
    # Add article-specific metadata if available
    if article_metadata is not None:
        for column in ['Authors', 'DOI', 'Publication Year', 'Journal']:
            if column in article_metadata and not pd.isna(article_metadata[column]):
                metadata_fields.append((column, article_metadata[column]))
    
    for field, value in metadata_fields:
        row_cells = metadata_table.add_row().cells
        row_cells[0].text = field
        row_cells[1].text = str(value)
    
    doc.add_paragraph()
    
    # Add each section
    section_order = [
        ("background", "Background"),
        ("methods", "Methods"),
        ("results", "Results"),
        ("discussion", "Discussion"),
        ("conclusions", "Conclusions"),
        ("synthesis", "Critical Synthesis")
    ]
    
    for section_key, section_title in section_order:
        if section_key in results.get("sections", {}):
            doc.add_heading(section_title, level=2)
            doc.add_paragraph(results["sections"][section_key])
            doc.add_paragraph()
    
    # Save the document if path provided
    if output_path:
        doc.save(output_path)
        print(f"Word document saved to: {output_path}")
    
    return doc

def main():
    """Main function to process PDF articles."""
    args = parse_arguments()
    base_dir, data_dir, articles_dir, output_dir, prompts_dir = setup_directories(args)
    
    # Initialize the prompt manager
    prompt_manager = PromptManager(args.prompts_dir if args.prompts_dir else prompts_dir)
    
    # Load article metadata
    article_metadata = load_article_metadata(data_dir)
    
    # Get PDF files to process
    pdf_files = get_pdf_files(articles_dir, args.article_id)
    
    if not pdf_files:
        print("No PDF files found to process.")
        return
    
    # Limit to a single file for testing if requested
    if args.test and len(pdf_files) > 1:
        pdf_files = pdf_files[:1]
        print(f"Test mode: Processing only the first PDF file: {pdf_files[0]}")
    
    # LLM model to use - you can modify this to use your preferred model
    model_name = args.model
    
    # Process each PDF
    for pdf_file in pdf_files:
        filename = os.path.basename(pdf_file)
        print(f"\nProcessing: {filename}")
        
        # Try to determine article ID from the filename
        article_id = None
        id_match = re.search(r'(\d+)[_\.]', filename)
        if id_match:
            article_id = int(id_match.group(1))
        
        # Extract article metadata if available
        article_meta = None
        if article_metadata is not None and article_id is not None:
            # Find matching article in metadata
            article_meta = article_metadata[article_metadata.index == article_id - 1].iloc[0] if article_id - 1 in article_metadata.index else None
            if article_meta is None:
                # Try searching by DOI if available
                if 'DOI' in article_metadata.columns:
                    doi_pattern = re.compile(r'10\.\d{4,9}/[-._;()/:A-Za-z0-9]+')
                    doi_match = doi_pattern.search(filename)
                    if doi_match:
                        doi = doi_match.group(0)
                        article_meta = article_metadata[article_metadata['DOI'].str.contains(doi, na=False)].iloc[0] if any(article_metadata['DOI'].str.contains(doi, na=False)) else None
        
        # Extract text from PDF using enhanced extraction
        print("Extracting text from PDF...")
        pdf_text = improve_pdf_extraction(pdf_file)
        
        if pdf_text is None:
            print(f"Skipping {filename} due to text extraction failure.")
            continue
        
        # Detect article structure to better inform processing
        article_structure = detect_article_structure(pdf_text)
        print(f"Detected article type: {article_structure['article_type']}")
        
        # Clean and prepare the text using enhanced cleaning
        print("Cleaning and preparing text...")
        clean_pdf_text = clean_scientific_text(pdf_text)
        
        if not clean_pdf_text:
            print(f"Skipping {filename} - no usable text after cleaning.")
            continue
        
        # Prepare article info for specialized prompts
        article_info = {}
        
        # Determine study design based on detected structure
        if article_structure['article_type'] == 'research':
            if article_structure.get('has_methods', False) and article_structure.get('has_results', False):
                article_info['study_design'] = 'quasi_experimental'  # Default to quasi-experimental unless RCT is detected
        elif article_structure['article_type'] == 'review':
            article_info['study_design'] = 'systematic_review'
        
        # If article metadata available, add more info
        if article_meta is not None:
            # Add mental health condition if available in metadata
            if 'Mental Health Condition' in article_meta:
                article_info['mental_health_condition'] = article_meta['Mental Health Condition']
            
            # Add intervention type if available in metadata
            if 'Intervention Type' in article_meta:
                article_info['intervention_type'] = article_meta['Intervention Type']
            
            # If study design is specified in metadata, use that
            if 'Study Design' in article_meta:
                article_info['study_design'] = article_meta['Study Design']
        
        # Process the article with LLM
        print(f"Processing article with {model_name}...")
        results = process_all_sections(model_name, clean_pdf_text, prompt_manager, article_info, article_id, pdf_file, args.api_url)
        
        # Add article structure information to results
        results["article_structure"] = article_structure
        
        # Create a unique output filename
        if article_id:
            output_filename = f"article_{article_id}_summary.docx"
        else:
            # Use the PDF filename but change extension to .docx
            output_filename = os.path.splitext(filename)[0] + "_summary.docx"
        
        output_path = os.path.join(output_dir, output_filename)
        
        # Create Word document
        print("Creating Word document...")
        create_word_document(results, article_meta, output_path)
        
        # Also save raw JSON results for reference
        json_path = os.path.join(output_dir, os.path.splitext(output_filename)[0] + ".json")
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(results, f, indent=2)
        
        print(f"Processing complete for {filename}")
        print(f"Outputs saved to:\n- {output_path}\n- {json_path}")
    
    print("\nAll articles processed successfully!")

if __name__ == "__main__":
    main()