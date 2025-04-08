#!/usr/bin/env python
"""
Batch PDF Article Processor

This script provides batch processing capabilities for analyzing multiple PDF articles.
It builds on the article_summarizer.py script to process multiple articles in parallel
and compile the results into a single comprehensive document.

Usage:
    python batch_processor.py [--model MODEL_NAME] [--parallel N] [--output OUTPUT_FILE]

Options:
    --model MODEL_NAME    Specify the LLM model to use (default: mistral-small-24b-instruct-2501)
    --parallel N          Number of articles to process in parallel (default: 2)
    --output OUTPUT_FILE  Name of the output file (default: combined_summary.docx)
    --prompts_dir DIR     Directory containing custom prompt templates
"""

import os
import sys
import argparse
import json
import pandas as pd
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import concurrent.futures
import time
from datetime import datetime
import tqdm
import logging

# Import from article_summarizer.py
sys.path.append(os.path.dirname(os.path.abspath(__file__)))
from article_summarizer import (
    setup_directories, load_article_metadata, get_pdf_files,
    process_all_sections, create_word_document
)

# Import PDF extraction utilities
from pdf_extraction_utils import (
    improve_pdf_extraction,
    clean_scientific_text,
    detect_article_structure
)

# Import the PromptManager for advanced prompt handling
from advanced_prompt_manager import PromptManager

# Set up logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("batch_processor.log"),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

def parse_arguments():
    """Parse command line arguments."""
    parser = argparse.ArgumentParser(description='Batch process multiple PDF articles using LLMs')
    parser.add_argument('--model', type=str, default='mistral-small-24b-instruct-2501',
                        help='LLM model to use for analysis')
    parser.add_argument('--parallel', type=int, default=2,
                        help='Number of articles to process in parallel')
    parser.add_argument('--output', type=str, default='combined_summary.docx',
                        help='Name of the output file')
    parser.add_argument('--prompts_dir', type=str, help='Directory containing custom prompt templates')
    parser.add_argument('--api_url', type=str, help='Custom LLM API URL')
    return parser.parse_args()

def process_article(args):
    """
    Process a single article with the given parameters.
    This function is designed to be used with concurrent.futures.
    
    Args:
        args (tuple): Tuple containing (pdf_path, model_name, prompt_manager, article_metadata, index, api_url)
        
    Returns:
        dict: Results of processing the article
    """
    pdf_path, model_name, prompt_manager, article_meta, index, api_url = args
    
    article_id = None
    if article_meta is not None:
        article_id = index + 1
    
    try:
        filename = os.path.basename(pdf_path)
        logger.info(f"Processing article {index+1}: {filename}")
        
        # Extract text from PDF using improved extraction
        pdf_text = improve_pdf_extraction(pdf_path)
        if pdf_text is None:
            logger.warning(f"Text extraction failed for {filename}")
            return None
        
        # Detect article structure
        article_structure = detect_article_structure(pdf_text)
        logger.info(f"Detected article type for {filename}: {article_structure['article_type']}")
        
        # Clean the text
        clean_pdf_text = clean_scientific_text(pdf_text)
        if not clean_pdf_text:
            logger.warning(f"No usable text after cleaning for {filename}")
            return None
        
        # Prepare article info for specialized prompts
        article_info = {}
        
        # Determine study design based on detected structure
        if article_structure['article_type'] == 'research':
            if article_structure.get('has_methods', False) and article_structure.get('has_results', False):
                article_info['study_design'] = 'quasi_experimental'  # Default unless RCT is detected
        elif article_structure['article_type'] == 'review':
            article_info['study_design'] = 'systematic_review'
        
        # If article metadata available, add more info
        if article_meta is not None:
            # Add mental health condition if available in metadata
            if 'Mental Health Condition' in article_meta:
                article_info['mental_health_condition'] = article_meta['Mental Health Condition']
            
            # Add intervention type if available in metadata
            if 'Intervention Type' in article_meta:
                article_info['intervention_type'] = article_meta['Intervention Type']
            
            # If study design is specified in metadata, use that
            if 'Study Design' in article_meta:
                article_info['study_design'] = article_meta['Study Design']
        
        # Process with LLM
        results = process_all_sections(model_name, clean_pdf_text, prompt_manager, article_info, article_id, pdf_path, api_url)
        
        # Add metadata
        if article_meta is not None:
            results["metadata"] = {
                "title": article_meta.get("Article Title", "Unknown Title"),
                "authors": article_meta.get("Authors", "Unknown Authors"),
                "doi": article_meta.get("DOI", "Unknown DOI"),
                "publication_year": article_meta.get("Publication Year", "Unknown Year"),
                "journal": article_meta.get("Journal", "Unknown Journal"),
                "mental_health_condition": article_meta.get("Mental Health Condition", "Unknown"),
                "intervention_type": article_meta.get("Intervention Type", "Unknown"),
                "study_design": article_meta.get("Study Design", article_info.get("study_design", "Unknown"))
            }
        
        # Add article structure information
        results["article_structure"] = article_structure
        
        return results
    
    except Exception as e:
        logger.error(f"Error processing article {index+1} ({pdf_path}): {str(e)}")
        return None

def create_combined_document(all_results, article_metadata, output_path, args, prompt_manager):
    """
    Create a single Word document containing summaries of all processed articles.
    
    Args:
        all_results (list): List of results from processing each article
        article_metadata (pd.DataFrame): DataFrame with article metadata
        output_path (str): Path to save the combined document
        args: Command line arguments containing model and API URL information
        prompt_manager (PromptManager): Instance of the prompt manager for getting prompts
        
    Returns:
        docx.Document: The created Word document
    """
    doc = docx.Document()
    
    # Set up document styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # Add title page
    title = doc.add_heading('Systematic Review: Article Summaries', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add date
    current_date = datetime.now().strftime("%B %d, %Y")
    date_paragraph = doc.add_paragraph(current_date)
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add table of contents
    doc.add_page_break()
    toc_heading = doc.add_heading('Table of Contents', level=1)
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add placeholder for table of contents
    doc.add_paragraph("(Table of contents will be populated manually after generation)")
    doc.add_page_break()
    
    # Generate executive summary using LLM
    doc.add_heading('Executive Summary', level=1)
    
    # Collect concise data about all articles to feed to the LLM
    articles_info = []
    for result in all_results:
        if result is None:
            continue
        
        article_title = result.get("metadata", {}).get("title", "Unknown Title")
        article_info = f"Article: {article_title}\n"
        
        if "sections" in result and "conclusions" in result["sections"]:
            article_info += f"Conclusions: {result['sections']['conclusions'][:300]}...\n"
        
        articles_info.append(article_info)
    
    combined_info = "\n\n".join(articles_info)
    executive_summary_prompt = prompt_manager.get_prompt("executive_summary")
    executive_summary = process_all_sections(args.model, combined_info, prompt_manager, None, None, None, args.api_url)
    
    if "sections" in executive_summary and "executive_summary" in executive_summary["sections"]:
        doc.add_paragraph(executive_summary["sections"]["executive_summary"])
    else:
        # Fallback if LLM processing fails
        doc.add_paragraph(
            "This document contains summaries of scientific articles related to employment "
            "interventions for people with mental health disorders. Each article has been analyzed "
            "to extract key information about the research background, methods, results, and conclusions. "
            "Additionally, a critical synthesis has been provided for each article to evaluate its "
            "contribution to the field."
        )
    
    doc.add_page_break()
    
    # Add each article summary
    for i, results in enumerate(all_results):
        if results is None:
            continue
        
        article_num = i + 1
        
        # Get article title
        if "metadata" in results and results["metadata"].get("title"):
            article_title = results["metadata"].get("title")
        elif results.get("filename"):
            article_title = f"Article {article_num}: {os.path.basename(results['filename'])}"
        else:
            article_title = f"Article {article_num}"
        
        # Add article heading
        doc.add_heading(f"{article_num}. {article_title}", level=1)
        
        # Add metadata section
        doc.add_heading("Article Metadata", level=2)
        
        metadata_table = doc.add_table(rows=1, cols=2)
        metadata_table.style = 'Table Grid'
        
        # Set column widths
        for cell in metadata_table.columns[0].cells:
            cell.width = Inches(1.5)
        for cell in metadata_table.columns[1].cells:
            cell.width = Inches(4.5)
        
        # Add metadata headers
        hdr_cells = metadata_table.rows[0].cells
        hdr_cells[0].text = "Field"
        hdr_cells[1].text = "Value"
        
        # Start with basic metadata
        metadata_fields = [
            ("Article ID", results.get("article_id", f"{article_num}")),
            ("Filename", os.path.basename(results.get("filename", "N/A"))),
        ]
        
        # Add article-specific metadata if available
        if "metadata" in results:
            for field, value in results["metadata"].items():
                if value and value != "Unknown":
                    # Format the field name for display
                    display_field = field.replace("_", " ").title()
                    metadata_fields.append((display_field, value))
        
        # Add metadata rows
        for field, value in metadata_fields:
            row_cells = metadata_table.add_row().cells
            row_cells[0].text = field
            row_cells[1].text = str(value)
        
        doc.add_paragraph()
        
        # Add each content section
        section_order = [
            ("background", "Background"),
            ("methods", "Methods"),
            ("results", "Results"),
            ("discussion", "Discussion"),
            ("conclusions", "Conclusions"),
            ("synthesis", "Critical Synthesis")
        ]
        
        for section_key, section_title in section_order:
            if section_key in results.get("sections", {}):
                doc.add_heading(section_title, level=2)
                doc.add_paragraph(results["sections"][section_key])
                doc.add_paragraph()
        
        # Add page break between articles
        if i < len(all_results) - 1:
            doc.add_page_break()
    
    # Generate overall analysis using LLM
    doc.add_page_break()
    doc.add_heading('Overall Analysis and Synthesis', level=1)
    
    # Collect key findings and synthesis sections from all articles
    synthesis_data = []
    for result in all_results:
        if result is None:
            continue
        
        if "sections" in result:
            article_synthesis = ""
            if "synthesis" in result["sections"]:
                article_synthesis += result["sections"]["synthesis"]
            elif "conclusions" in result["sections"]:
                article_synthesis += result["sections"]["conclusions"]
            
            synthesis_data.append(article_synthesis)
    
    combined_synthesis = "\n\n".join(synthesis_data)
    overall_synthesis_prompt = prompt_manager.get_prompt("overall_synthesis")
    overall_synthesis = process_all_sections(args.model, combined_synthesis, prompt_manager, None, None, None, args.api_url)
    
    if "sections" in overall_synthesis and "overall_synthesis" in overall_synthesis["sections"]:
        doc.add_paragraph(overall_synthesis["sections"]["overall_synthesis"])
    else:
        # Fallback if LLM processing fails
        doc.add_paragraph(
            "This section provides an integrated analysis across all articles reviewed. "
            "It identifies patterns, consensus findings, and areas of disagreement in the literature. "
            "(To be completed manually after reviewing all article summaries.)"
        )
    
    # Save the document
    doc.save(output_path)
    logger.info(f"Combined document saved to: {output_path}")
    
    return doc

def main():
    """Main function to batch process multiple articles."""
    args = parse_arguments()
    base_dir, data_dir, articles_dir, output_dir, prompts_dir = setup_directories(args)
    
    # Initialize the prompt manager
    prompt_manager = PromptManager(args.prompts_dir if args.prompts_dir else prompts_dir)
    
    # Load article metadata
    article_metadata = load_article_metadata(data_dir)
    
    # Get all PDF files
    pdf_files = get_pdf_files(articles_dir)
    
    if not pdf_files:
        logger.error("No PDF files found to process.")
        return
    
    logger.info(f"Found {len(pdf_files)} PDF files to process")
    
    # Prepare processing arguments
    processing_args = []
    for i, pdf_file in enumerate(pdf_files):
        # Extract article metadata if available
        article_meta = None
        if article_metadata is not None and i < len(article_metadata):
            article_meta = article_metadata.iloc[i]
        
        # Include prompt_manager and api_url in args
        processing_args.append((pdf_file, args.model, prompt_manager, article_meta, i, args.api_url))
    
    # Process articles in parallel
    all_results = []
    with concurrent.futures.ProcessPoolExecutor(max_workers=args.parallel) as executor:
        futures = {executor.submit(process_article, arg): arg for arg in processing_args}
        
        # Use tqdm for progress bar
        with tqdm.tqdm(total=len(futures), desc="Processing articles") as progress:
            for future in concurrent.futures.as_completed(futures):
                arg = futures[future]
                try:
                    result = future.result()
                    if result:
                        all_results.append(result)
                    progress.update(1)
                except Exception as e:
                    pdf_path = arg[0]
                    logger.error(f"Error processing {os.path.basename(pdf_path)}: {str(e)}")
                    progress.update(1)
    
    # Save individual results
    logger.info(f"Successfully processed {len(all_results)} articles")
    for i, result in enumerate(all_results):
        # Create a unique output filename
        article_id = result.get("article_id")
        if article_id and article_id != "unknown":
            output_filename = f"article_{article_id}_summary.docx"
        else:
            filename = os.path.basename(result.get("filename", f"article_{i+1}"))
            output_filename = os.path.splitext(filename)[0] + "_summary.docx"
        
        output_path = os.path.join(output_dir, output_filename)
        
        # Create individual Word document
        article_meta = None
        if article_metadata is not None and i < len(article_metadata):
            article_meta = article_metadata.iloc[i]
        
        create_word_document(result, article_meta, output_path)
        
        # Save raw JSON results
        json_path = os.path.join(output_dir, os.path.splitext(output_filename)[0] + ".json")
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(result, f, indent=2)
    
    # Create combined document
    combined_output_path = os.path.join(output_dir, args.output)
    create_combined_document(all_results, article_metadata, combined_output_path, args, prompt_manager)
    
    logger.info("All processing complete!")
    logger.info(f"Combined document saved to: {combined_output_path}")
    logger.info(f"Individual article summaries saved to: {output_dir}")

if __name__ == "__main__":
    main()